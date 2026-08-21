#!/usr/bin/env python3
"""Generate a Word (.docx) document from text input."""
import sys, os, re

def _add_formatted_text(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|__.*?__)', text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        if re.match(r'^\*\*.*\*\*$', part):
            run.text = part[2:-2]; run.bold = True
        elif re.match(r'^\*.*\*$', part):
            run.text = part[1:-1]; run.italic = True
        elif re.match(r'^__.*__$', part):
            run.text = part[2:-2]; run.underline = True
        else:
            run.text = part
    return paragraph

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if len(sys.argv) < 3:
        print("Usage: python3 generate_docx.py <input_text_file> <output_path> [title]", file=sys.stderr)
        sys.exit(1)

    text_file, output_path = sys.argv[1], sys.argv[2]
    title = sys.argv[3] if len(sys.argv) > 3 else "Document"

    with open(text_file, 'r', encoding='utf-8') as f:
        text = f.read()

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'; style.font.size = Pt(11)
    doc.add_heading(title, level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    lines = text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line: i += 1; continue
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1); i += 1; continue
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2); i += 1; continue
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3); i += 1; continue
        bullet_match = re.match(r'^[\-\*\u2022]\s+', line)
        if bullet_match:
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, line[bullet_match.end():]); i += 1; continue
        numbered_match = re.match(r'^(\d+[\.\)]\s+)', line)
        if numbered_match:
            p = doc.add_paragraph(style='List Number')
            _add_formatted_text(p, line[numbered_match.end():]); i += 1; continue
        p = doc.add_paragraph()
        _add_formatted_text(p, line); i += 1

    doc.save(output_path)
    print(output_path)
except ImportError as e:
    print(f"Error importing dependencies: {e}", file=sys.stderr); sys.exit(1)
except Exception as e:
    print(f"Error generating Word document: {e}", file=sys.stderr); sys.exit(1)
